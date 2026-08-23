import io
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def format_sources_plain(sources):
    """Formats source documents into a plain list of strings."""
    lines = []
    if not sources:
        return lines
    seen = set()
    for doc in sources:
        if isinstance(doc, dict):
            fname = doc.get("source_file", "Document")
            page = doc.get("page", "?")
        elif hasattr(doc, "metadata"):
            fname = doc.metadata.get("source_file", "Document")
            page = doc.metadata.get("page", "?")
        else:
            continue
        key = (fname, page)
        if key not in seen:
            lines.append(f"- {fname}, Page {page}")
            seen.add(key)
    return lines


def format_verification_plain(verification):
    """Formats verification items into a plain list of strings."""
    lines = []
    if not verification:
        return lines
    for item in verification:
        if isinstance(item, dict):
            status = "Supported" if item.get("supported", False) else "Needs Review"
            sim = item.get("similarity", 0.0)
            sent = item.get("sentence", "")
            lines.append(f"- [{status}] ({sim:.2f}) {sent}")
    return lines


# ============================================================
# 1. TEXT EXPORT
# ============================================================
def export_turn_to_txt(question: str, answer: str, sources: list, verification: list, timestamp: str = "") -> str:
    """Generates plain text export for a single Q&A turn."""
    ts = timestamp or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sources_lines = format_sources_plain(sources)
    verif_lines = format_verification_plain(verification)
    
    txt = (
        f"================================================================================\n"
        f"CHAT WITH YOUR PDF — VERIFIED Q&A REPORT\n"
        f"Timestamp: {ts}\n"
        f"================================================================================\n\n"
        f"QUESTION:\n{question}\n\n"
        f"ANSWER:\n{answer}\n\n"
        f"RETRIEVED SOURCES:\n" + ("\n".join(sources_lines) if sources_lines else "None") + "\n\n"
        f"GROUNDING & VERIFICATION REPORT:\n" + ("\n".join(verif_lines) if verif_lines else "No verification data") + "\n\n"
        f"================================================================================\n"
    )
    return txt


def export_conversation_to_txt(username: str, history: list) -> str:
    """Generates plain text export for the full chat history."""
    txt = (
        f"================================================================================\n"
        f"CHAT WITH YOUR PDF — FULL CONVERSATION EXPORT\n"
        f"User: {username}\n"
        f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        f"Total Exchanges: {len(history)}\n"
        f"================================================================================\n\n"
    )
    for idx, turn in enumerate(history, 1):
        txt += f"--- QUESTION {idx} ({turn.get('timestamp', '')}) ---\n"
        txt += f"Q: {turn['question']}\n\n"
        txt += f"A: {turn['answer']}\n\n"
        sources_lines = format_sources_plain(turn.get("sources", []))
        txt += "Sources:\n" + ("\n".join(sources_lines) if sources_lines else " - None") + "\n\n"
        verif_lines = format_verification_plain(turn.get("verification", []))
        txt += "Verification:\n" + ("\n".join(verif_lines) if verif_lines else " - No data") + "\n\n"
        txt += "-" * 60 + "\n\n"
    return txt


# ============================================================
# 2. PDF EXPORT (ReportLab In-Memory)
# ============================================================
def export_turn_to_pdf(question: str, answer: str, sources: list, verification: list, timestamp: str = "") -> bytes:
    """Generates styled in-memory PDF for a single Q&A turn."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    primary_color = colors.HexColor("#1E1B4B")
    accent_color = colors.HexColor("#FF2E93")
    text_dark = colors.HexColor("#1F2937")
    
    title_style = ParagraphStyle("T1", fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=primary_color, spaceAfter=4)
    sub_style = ParagraphStyle("Sub", fontName="Helvetica", fontSize=10, leading=14, textColor=colors.HexColor("#6B7280"), spaceAfter=12)
    h2_style = ParagraphStyle("H2", fontName="Helvetica-Bold", fontSize=12, leading=16, textColor=primary_color, spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle("B", fontName="Helvetica", fontSize=10, leading=14, textColor=text_dark, spaceAfter=6)
    q_style = ParagraphStyle("Q", fontName="Helvetica-Bold", fontSize=11, leading=15, textColor=colors.HexColor("#0F172A"), spaceAfter=6)
    code_style = ParagraphStyle("C", fontName="Helvetica", fontSize=9, leading=13, textColor=colors.HexColor("#334155"), leftIndent=10, spaceAfter=3)

    story = []
    
    story.append(Paragraph("Chat With Your PDF · Verified Answer Report", title_style))
    ts = timestamp or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    story.append(Paragraph(f"Generated on {ts} · Grounded Document Intelligence", sub_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=accent_color, spaceBefore=2, spaceAfter=14))
    
    story.append(Paragraph("QUESTION", h2_style))
    story.append(Paragraph(question.replace("<", "&lt;").replace(">", "&gt;"), q_style))
    story.append(Spacer(1, 6))
    
    story.append(Paragraph("AI-GENERATED ANSWER", h2_style))
    clean_ans = answer.replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    story.append(Paragraph(clean_ans, body_style))
    story.append(Spacer(1, 8))
    
    sources_lines = format_sources_plain(sources)
    if sources_lines:
        story.append(Paragraph("RETRIEVED SOURCES & CITATIONS", h2_style))
        for s in sources_lines:
            story.append(Paragraph(f"<b>&bull;</b> {s.replace('- ', '')}", code_style))
        story.append(Spacer(1, 8))
        
    if verification:
        story.append(Paragraph("GROUNDING & VERIFICATION AUDIT", h2_style))
        table_rows = [
            [Paragraph("<b>Status</b>", body_style), Paragraph("<b>Score</b>", body_style), Paragraph("<b>Sentence</b>", body_style)]
        ]
        for item in verification:
            if isinstance(item, dict):
                is_sup = item.get("supported", False)
                sim = item.get("similarity", 0.0)
                sent = item.get("sentence", "")
                status_p = Paragraph("<font color='#059669'><b>Supported</b></font>" if is_sup else "<font color='#D97706'><b>Review</b></font>", body_style)
                score_p = Paragraph(f"<b>{sim:.2f}</b>", body_style)
                sent_p = Paragraph(sent.replace("<", "&lt;").replace(">", "&gt;"), code_style)
                table_rows.append([status_p, score_p, sent_p])
                
        t_v = Table(table_rows, colWidths=[75, 45, 410])
        t_v.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#EDE9FE")),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E1")),
            ('PADDING', (0,0), (-1,-1), 4),
            ('VALIGN', (0,0), (-1,-1), 'TOP')
        ]))
        story.append(t_v)
        
    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


def export_conversation_to_pdf(username: str, history: list) -> bytes:
    """Generates styled in-memory PDF for full chat history."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    primary_color = colors.HexColor("#1E1B4B")
    accent_color = colors.HexColor("#FF2E93")
    
    title_style = ParagraphStyle("T1", fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=primary_color, spaceAfter=4)
    sub_style = ParagraphStyle("Sub", fontName="Helvetica", fontSize=10, leading=14, textColor=colors.HexColor("#6B7280"), spaceAfter=12)
    h2_style = ParagraphStyle("H2", fontName="Helvetica-Bold", fontSize=12, leading=16, textColor=primary_color, spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle("B", fontName="Helvetica", fontSize=9.5, leading=13.5, textColor=colors.HexColor("#1F2937"), spaceAfter=4)
    q_style = ParagraphStyle("Q", fontName="Helvetica-Bold", fontSize=10.5, leading=14.5, textColor=colors.HexColor("#0F172A"), spaceAfter=4)
    code_style = ParagraphStyle("C", fontName="Helvetica", fontSize=8.5, leading=11.5, textColor=colors.HexColor("#475569"), leftIndent=10, spaceAfter=2)

    story = []
    
    story.append(Paragraph("Chat With Your PDF · Full Conversation Export", title_style))
    story.append(Paragraph(f"User: <b>{username}</b> · Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} · Total Turns: {len(history)}", sub_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=accent_color, spaceBefore=2, spaceAfter=14))
    
    for idx, turn in enumerate(history, 1):
        story.append(Paragraph(f"Turn #{idx} — {turn.get('timestamp', '')}", h2_style))
        story.append(Paragraph(f"<b>Q:</b> {turn['question'].replace('<', '&lt;').replace('>', '&gt;')}", q_style))
        clean_ans = turn['answer'].replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
        story.append(Paragraph(f"<b>A:</b> {clean_ans}", body_style))
        
        sources_lines = format_sources_plain(turn.get("sources", []))
        if sources_lines:
            story.append(Paragraph("<b>Sources:</b> " + " | ".join([s.replace("- ", "") for s in sources_lines]), code_style))
            
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#E2E8F0"), spaceBefore=6, spaceAfter=10))
        
    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


# ============================================================
# 3. DOCX EXPORT (python-docx In-Memory)
# ============================================================
def export_turn_to_docx(question: str, answer: str, sources: list, verification: list, timestamp: str = "") -> bytes:
    """Generates styled in-memory DOCX (Word) document for a single Q&A turn."""
    doc = Document()
    
    p_title = doc.add_heading("Chat With Your PDF · Q&A Report", level=0)
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    ts = timestamp or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    p_sub = doc.add_paragraph(f"Export Date: {ts} | Grounded Document Intelligence")
    p_sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    p_sub.runs[0].font.size = Pt(10)
    
    doc.add_heading("Question", level=1)
    p_q = doc.add_paragraph(question)
    p_q.runs[0].font.bold = True
    p_q.runs[0].font.size = Pt(11)
    
    doc.add_heading("Answer", level=1)
    doc.add_paragraph(answer)
    
    sources_lines = format_sources_plain(sources)
    if sources_lines:
        doc.add_heading("Retrieved Sources", level=2)
        for s in sources_lines:
            doc.add_paragraph(s.replace("- ", ""), style='List Bullet')
            
    if verification:
        doc.add_heading("Grounding & Verification Audit", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Status"
        hdr_cells[1].text = "Score"
        hdr_cells[2].text = "Sentence"
        
        for item in verification:
            if isinstance(item, dict):
                row_cells = table.add_row().cells
                is_sup = item.get("supported", False)
                sim = item.get("similarity", 0.0)
                sent = item.get("sentence", "")
                row_cells[0].text = "Supported" if is_sup else "Needs Review"
                row_cells[1].text = f"{sim:.2f}"
                row_cells[2].text = sent
                
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes


def export_conversation_to_docx(username: str, history: list) -> bytes:
    """Generates styled in-memory DOCX document for the full chat history."""
    doc = Document()
    
    p_title = doc.add_heading("Chat With Your PDF — Full Conversation", level=0)
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p_sub = doc.add_paragraph(f"User: {username} | Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Total Turns: {len(history)}")
    p_sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    
    for idx, turn in enumerate(history, 1):
        doc.add_heading(f"Turn #{idx} ({turn.get('timestamp', '')})", level=1)
        
        p_q = doc.add_paragraph()
        p_q.add_run("Question: ").bold = True
        p_q.add_run(turn['question'])
        
        p_a = doc.add_paragraph()
        p_a.add_run("Answer: ").bold = True
        p_a.add_run(turn['answer'])
        
        sources_lines = format_sources_plain(turn.get("sources", []))
        if sources_lines:
            p_s = doc.add_paragraph()
            p_s.add_run("Sources: ").bold = True
            p_s.add_run(" | ".join([s.replace("- ", "") for s in sources_lines]))
            
        doc.add_paragraph("-" * 40)
        
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes
